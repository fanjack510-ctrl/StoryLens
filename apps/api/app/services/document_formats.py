"""从各种文件格式里读出「大纲」——能直接读到的就别猜。

原来的摄入把所有格式先拍平成一个大字符串，再拿小说的「第 N 章」正则去切。对小说 txt 没问题，
对专著是把现成的东西扔了：Word 的标题样式、Markdown 的 `##`、LaTeX 的 `\\section`、HTML 的
`<h2>`、EPUB 的目录——**这些格式里层级是明写的**，读一下就有，比从纯文本猜准得多，也解释得清。

所以这里分成两支：

    有结构的格式   docx / md / tex / html / epub / odt   直接读，source="declared"
    没结构的格式   pdf / txt                              走推断，source="inferred"

界面要能说出结构是读来的还是猜的。读来的可以直接信；猜的必须让用户看见规则和覆盖率——
一份缺了节的摘要，读者不会知道自己漏了什么。

中文书和英文书都要认。中文专著的编号习惯完全不同（`第三章` / `一、` / `（二）` / `§2`），
而且常常没有英文手册那种「章首自带带页码的小节目录」。

**不支持的要明说。** 知网的 CAJ/NH/KDH 是专有格式，没有可用的开源解析器；扫描版 PDF 需要先
OCR。这两类都单独抛错——含糊的一句「无法导入」会让用户反复换文件，而不是去做该做的那件事。
"""

from __future__ import annotations

import re
import zipfile
from io import BytesIO
from pathlib import Path

from app.domain.document_outline import BookOutline, OutlineNode, outline_from_headings

__all__ = [
    "SUPPORTED_SUFFIXES",
    "UnsupportedFormatError",
    "outline_from_bytes",
    "outline_from_docx",
    "outline_from_markdown",
    "outline_from_latex",
    "outline_from_html",
    "outline_from_odt",
]


class UnsupportedFormatError(ValueError):
    """这个格式打不开，而且要说清楚为什么、以及该怎么办。"""


#: 明确不支持，且原因各不相同——错误信息必须分开说。
_KNOWN_UNSUPPORTED = {
    ".caj": "知网 CAJ 是专有格式，没有可用的解析器。请在知网页面选择「PDF 下载」后再导入。",
    ".nh": "知网 NH 是专有格式，请改用 PDF 下载。",
    ".kdh": "知网 KDH 是专有格式，请改用 PDF 下载。",
    ".mobi": "MOBI 暂不支持，请用 Calibre 转成 EPUB 后导入。",
    ".azw3": "AZW3 暂不支持，请用 Calibre 转成 EPUB 后导入。",
    ".rtf": "RTF 暂不支持，请另存为 DOCX 后导入。",
    ".djvu": "DjVu 是扫描格式，需要先 OCR 成有文字层的 PDF。",
}

SUPPORTED_SUFFIXES = (
    ".txt", ".docx", ".epub", ".pdf",
    ".md", ".markdown", ".tex", ".latex",
    ".html", ".htm", ".xhtml", ".odt",
)


# --------------------------------------------------------------------- Word

#: Word 的标题样式：英文 "Heading 1"，中文 "标题 1"，还有 Title/副标题
_DOCX_HEADING = re.compile(r"^(?:Heading|标题)\s*(\d)", re.I)


def outline_from_docx(content: bytes) -> BookOutline:
    """Word 的标题层级来自段落样式，不是靠字号猜的。

    以前只取 `paragraph.text`，等于把整份文档的结构信息扔掉，再从纯文本里猜回来。
    """
    from docx import Document

    doc = Document(BytesIO(content))
    outline = BookOutline(source="declared")
    outline.rules.append("Word 标题样式（Heading n / 标题 n）")
    current: OutlineNode | None = None
    front: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = getattr(getattr(para, "style", None), "name", "") or ""
        m = _DOCX_HEADING.match(style)
        if m or style.lower() in {"title", "标题"}:
            level = int(m.group(1)) if m else 1
            current = OutlineNode(level=level, number="", title=text)
            outline.nodes.append(current)
            continue
        if current is None:
            front.append(text)
        else:
            current.paragraphs.append(text)
    if front:
        outline.nodes.insert(
            0, OutlineNode(level=1, number="", title="前置内容", paragraphs=front)
        )
    if not outline.nodes:
        return BookOutline(source="inferred")
    return outline


# ----------------------------------------------------------------- Markdown

_MD_HEADING = re.compile(r"^(#{1,6})\s+(\S.*)$")


def outline_from_markdown(text: str) -> BookOutline:
    outline = BookOutline(source="declared")
    outline.rules.append("Markdown 标题（#）")
    current: OutlineNode | None = None
    front: list[str] = []
    for raw in text.splitlines():
        m = _MD_HEADING.match(raw.strip())
        if m:
            current = OutlineNode(level=len(m.group(1)), number="", title=m.group(2).strip())
            outline.nodes.append(current)
            continue
        line = raw.strip()
        if not line:
            continue
        (current.paragraphs if current else front).append(line)
    if front:
        outline.nodes.insert(0, OutlineNode(1, "", "前置内容", front))
    return outline if outline.nodes else BookOutline(source="inferred")


# -------------------------------------------------------------------- LaTeX

_TEX_LEVEL = {"part": 1, "chapter": 1, "section": 2, "subsection": 3, "subsubsection": 4}
_TEX_HEADING = re.compile(r"\\(part|chapter|section|subsection|subsubsection)\*?\s*\{")
_TEX_COMMENT = re.compile(r"(?<!\\)%.*$", re.M)


def _tex_balanced(text: str, start: int) -> tuple[str, int]:
    """取出 `{...}` 里的内容——标题里常有嵌套花括号（`\\emph{...}`），不能用非贪婪正则。"""
    depth, out, i = 0, [], start
    while i < len(text):
        ch = text[i]
        if ch == "{":
            depth += 1
            if depth == 1:
                i += 1
                continue
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return "".join(out), i + 1
        if depth >= 1:
            out.append(ch)
        i += 1
    return "".join(out), len(text)


def _tex_plain(text: str) -> str:
    text = re.sub(r"\\[a-zA-Z]+\*?\s*(\[[^\]]*\])?", " ", text)
    return re.sub(r"[{}]", "", text)


def outline_from_latex(text: str) -> BookOutline:
    """论文最常见的源格式。层级由 `\\section` 一族直接给出，比任何推断都准。"""
    body = _TEX_COMMENT.sub("", text)
    m = re.search(r"\\begin\{document\}", body)
    if m:
        body = body[m.end():]
    outline = BookOutline(source="declared")
    outline.rules.append("LaTeX 章节命令（\\chapter / \\section / …）")
    pos, current, front = 0, None, []
    for m in _TEX_HEADING.finditer(body):
        chunk = body[pos:m.start()]
        text_chunk = [x.strip() for x in _tex_plain(chunk).splitlines() if x.strip()]
        (current.paragraphs if current else front).extend(text_chunk)
        title, after = _tex_balanced(body, m.end() - 1)
        current = OutlineNode(
            level=_TEX_LEVEL[m.group(1)], number="", title=_tex_plain(title).strip()
        )
        outline.nodes.append(current)
        pos = after
    tail = [x.strip() for x in _tex_plain(body[pos:]).splitlines() if x.strip()]
    (current.paragraphs if current else front).extend(tail)
    if front:
        outline.nodes.insert(0, OutlineNode(1, "", "前置内容", front))
    return outline if outline.nodes else BookOutline(source="inferred")


# --------------------------------------------------------------------- HTML


def outline_from_html(markup: str) -> BookOutline:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(markup, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    outline = BookOutline(source="declared")
    outline.rules.append("HTML 标题标签（h1–h6）")
    current: OutlineNode | None = None
    front: list[str] = []
    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "div"]):
        name = el.name
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        if name and name.startswith("h") and len(name) == 2 and name[1].isdigit():
            current = OutlineNode(level=int(name[1]), number="", title=text)
            outline.nodes.append(current)
            continue
        if name == "div" and el.find(["p", "li", "div", "h1", "h2", "h3"]):
            continue          # 容器 div 的文字会和子元素重复
        (current.paragraphs if current else front).append(text)
    if front:
        outline.nodes.insert(0, OutlineNode(1, "", "前置内容", front))
    return outline if outline.nodes else BookOutline(source="inferred")


# ---------------------------------------------------------------------- ODT


def outline_from_odt(content: bytes) -> BookOutline:
    """ODT 是一个 zip，结构在 content.xml 里。用 lxml 直接读，不为它引一个新依赖。"""
    from lxml import etree

    with zipfile.ZipFile(BytesIO(content)) as zf:
        xml = zf.read("content.xml")
    root = etree.fromstring(xml)
    ns = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }
    outline = BookOutline(source="declared")
    outline.rules.append("ODT 大纲级别（text:outline-level）")
    current: OutlineNode | None = None
    front: list[str] = []
    for el in root.iter():
        tag = etree.QName(el).localname
        if tag not in {"h", "p"}:
            continue
        text = "".join(el.itertext()).strip()
        if not text:
            continue
        if tag == "h":
            level = el.get(f"{{{ns['text']}}}outline-level") or "1"
            current = OutlineNode(level=int(level), number="", title=text)
            outline.nodes.append(current)
        else:
            (current.paragraphs if current else front).append(text)
    if front:
        outline.nodes.insert(0, OutlineNode(1, "", "前置内容", front))
    return outline if outline.nodes else BookOutline(source="inferred")


# ------------------------------------------------------------------ dispatch


def outline_from_bytes(filename: str, content: bytes) -> BookOutline:
    """按后缀选路。结构读得到就读，读不到才猜。"""
    suffix = Path(filename).suffix.lower()
    if suffix in _KNOWN_UNSUPPORTED:
        raise UnsupportedFormatError(_KNOWN_UNSUPPORTED[suffix])

    if suffix == ".docx":
        return outline_from_docx(content)
    if suffix in {".md", ".markdown"}:
        return outline_from_markdown(content.decode("utf-8", errors="replace"))
    if suffix in {".tex", ".latex"}:
        return outline_from_latex(content.decode("utf-8", errors="replace"))
    if suffix in {".html", ".htm", ".xhtml"}:
        return outline_from_html(content.decode("utf-8", errors="replace"))
    if suffix == ".odt":
        return outline_from_odt(content)

    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedFormatError(
            "暂不支持这个格式。目前支持：" + "、".join(s.lstrip(".").upper() for s in SUPPORTED_SUFFIXES)
        )

    # pdf / txt / epub：没有可直接读取的层级，交给推断
    from app.services.extractors import extract_document

    doc = extract_document(filename, content)

    if doc.pages:
        # PDF 有页边界，就走专门那条：章首目录 + 抹空白匹配。逐行认标题在这类版面上会把目录
        # 行本身当成标题（实测 264 节、71 个没有正文），而按页那条是 76 节、零遗漏。
        # 页边界是 PDF 独有的信息，不用等于白扔。
        from app.domain.monograph_ingestion import detect_monograph

        det = detect_monograph(doc.pages)
        if det.sections:
            outline = BookOutline(source="inferred")
            outline.rules.extend(det.rules)
            for sec in det.sections:
                outline.nodes.append(
                    OutlineNode(
                        level=min(4, sec.number.count(".") + 1),
                        number=sec.number,
                        title=sec.title,
                        paragraphs=list(sec.paragraphs),
                        chapter=f"第{sec.chapter_no}章" if sec.chapter_no is not None else "",
                    )
                )
            return outline

    return outline_from_headings(doc.text.splitlines(), source="inferred")
