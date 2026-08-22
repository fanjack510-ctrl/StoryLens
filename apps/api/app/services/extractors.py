from io import BytesIO
from pathlib import Path
from dataclasses import dataclass

from bs4 import BeautifulSoup
from docx import Document
from ebooklib import ITEM_DOCUMENT, epub


class InvalidFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class ScannedDocumentError(ValueError):
    """PDF 里没有文字层——是一沓图片，不是一本书。

    这一条必须单独抛，不能并进 EmptyDocumentError：两者对用户的意义完全不同。「文件里没有
    可导入的文本」听起来像文件坏了；而扫描件是好文件，只是还需要一步 OCR。把它们混成一句，
    用户会去换文件，而不是去做该做的那件事。
    """


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    encoding: str
    bom: str
    newline: str
    byte_count: int
    #: PDF 的逐页文本。专著的结构识别要靠页边界（页眉重复、章首目录），拼成一整个字符串
    #: 之后这些信息就没了。其他格式为 None。
    pages: tuple[str, ...] | None = None


def extract_text(filename: str, content: bytes) -> str:
    return extract_document(filename, content).text


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        text, encoding, bom = _decode_txt(content)
    elif suffix == ".docx":
        text = "\n".join(paragraph.text for paragraph in Document(BytesIO(content)).paragraphs)
        encoding, bom = "docx/xml", "none"
    elif suffix == ".epub":
        text = _extract_epub(content)
        encoding, bom = "epub/html", "none"
    elif suffix == ".pdf":
        pages = _extract_pdf_pages(content)
        text = "\n".join(pages)
        if not text.strip():
            raise ScannedDocumentError(
                "这个 PDF 没有文字层（整本是扫描图片），需要先做 OCR 才能导入。"
            )
        newline = "CRLF" if "\r\n" in text else ("CR" if "\r" in text else "LF")
        return ExtractedDocument(text, "pdf/text", "none", newline, len(content), tuple(pages))
    else:
        raise InvalidFileTypeError("仅支持 TXT、DOCX、EPUB、PDF")
    if not text.strip():
        raise EmptyDocumentError("文件中没有可导入的文本")
    newline = "CRLF" if "\r\n" in text else ("CR" if "\r" in text else "LF")
    return ExtractedDocument(text, encoding, bom, newline, len(content))


def _extract_pdf_pages(content: bytes) -> list[str]:
    """逐页取文字。页边界要留着——专著的结构就藏在页眉和章首目录里，
    拼成一整个字符串之后就再也分不出来了。"""
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:  # noqa: BLE001 — pypdf 的异常层次很杂，一律翻译成人话
        # 后缀是 .pdf 不等于内容是 PDF。不接住这里，用户拿到的是
        # `PdfStreamError: Stream has ended unexpectedly` —— 一句他无从下手的话。
        raise InvalidFileTypeError("这个文件不是有效的 PDF，或者已经损坏") from exc
    out: list[str] = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — 单页解析失败不该让整本书都导不进来
            out.append("")
    return out


def _decode_txt(content: bytes) -> tuple[str, str, str]:
    bom_candidates = (
        (b"\xef\xbb\xbf", "utf-8-sig", "UTF-8"),
        (b"\xff\xfe", "utf-16-le", "UTF-16-LE"),
        (b"\xfe\xff", "utf-16-be", "UTF-16-BE"),
    )
    for marker, encoding, label in bom_candidates:
        if content.startswith(marker):
            text = content.decode(encoding)
            return text.lstrip("\ufeff"), encoding, label
    if b"\x00" in content[:4096]:
        for encoding in ("utf-16-le", "utf-16-be"):
            try:
                text = content.decode(encoding)
                if text.count("\ufffd") == 0:
                    return text.lstrip("\ufeff"), encoding, "none"
            except UnicodeDecodeError:
                continue
    for encoding in ("utf-8", "gb18030"):
        try:
            text = content.decode(encoding)
            if "\ufffd" in text:
                continue
            return text, encoding, "none"
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 文件编码无法可靠识别，请转换为 UTF-8、GB18030 或 UTF-16")


def _extract_epub(content: bytes) -> str:
    book = epub.read_epub(BytesIO(content))
    sections: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        sections.append(soup.get_text("\n"))
    return "\n".join(sections)
