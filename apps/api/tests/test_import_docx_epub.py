"""STEP 2.6 — DOCX/EPUB import regression + original file protection.

Offline only: FakeProvider via conftest client. No Live API.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from ebooklib import epub
from fastapi.testclient import TestClient
from sqlalchemy import func, select

from app.services.extractors import extract_document


def _minimal_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("第一章 试炼", level=1)
    doc.add_paragraph("林澈站在港口，手里攥着旧信。")
    doc.add_paragraph("风把浪推到堤岸上。")
    doc.add_heading("第二章 结局", level=1)
    doc.add_paragraph("他爬上灯塔顶层，点燃火种。")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _minimal_epub_bytes() -> bytes:
    book = epub.EpubBook()
    book.set_identifier("storylens-step26-epub")
    book.set_title("灯塔试炼 EPUB")
    book.set_language("zh")
    c1 = epub.EpubHtml(title="第一章", file_name="chap1.xhtml", lang="zh")
    c1.content = (
        "<h1>第一章 启程</h1>"
        "<p>林澈站在港口，手里攥着旧信。</p>"
        "<p>风把浪推到堤岸上。</p>"
    )
    c2 = epub.EpubHtml(title="第二章", file_name="chap2.xhtml", lang="zh")
    c2.content = "<h1>第二章 结局</h1><p>他爬上灯塔顶层，点燃火种。</p>"
    book.add_item(c1)
    book.add_item(c2)
    book.toc = (c1, c2)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", c1, c2]
    buf = BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


def test_extract_docx_and_epub_offline():
    docx_bytes = _minimal_docx_bytes()
    epub_bytes = _minimal_epub_bytes()
    d = extract_document("sample.docx", docx_bytes)
    e = extract_document("sample.epub", epub_bytes)
    assert "林澈" in d.text
    assert "灯塔" in d.text or "火种" in d.text
    assert "林澈" in e.text
    assert len(d.text.strip()) > 0
    assert len(e.text.strip()) > 0


def test_docx_import_api_and_original_file_kept(client: TestClient, tmp_path: Path):
    fixture = tmp_path / "import_sample.docx"
    payload = _minimal_docx_bytes()
    fixture.write_bytes(payload)
    assert fixture.is_file()

    response = client.post(
        "/api/v1/books/import",
        files={
            "file": (
                "import_sample.docx",
                payload,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201, response.text
    body = response.json()
    book_id = int(body["book_id"])
    assert body["chapter_count"] >= 1
    assert body["paragraph_count"] >= 1

    chapters = client.get(f"/api/v1/books/{book_id}/chapters").json()
    assert chapters
    paragraphs = client.get(f"/api/v1/chapters/{chapters[0]['id']}/paragraphs").json()
    assert paragraphs
    assert paragraphs[0]["id"].startswith(f"B{book_id:04d}-")

    deleted = client.delete(f"/api/v1/books/{book_id}")
    assert deleted.status_code in {200, 204}, deleted.text
    assert client.get(f"/api/v1/books/{book_id}").status_code == 404
    assert fixture.is_file()
    assert fixture.read_bytes() == payload


def test_epub_import_api_and_original_file_kept(client: TestClient, tmp_path: Path):
    fixture = tmp_path / "import_sample.epub"
    payload = _minimal_epub_bytes()
    fixture.write_bytes(payload)
    assert fixture.is_file()

    response = client.post(
        "/api/v1/books/import",
        files={"file": ("import_sample.epub", payload, "application/epub+zip")},
    )
    assert response.status_code == 201, response.text
    body = response.json()
    book_id = int(body["book_id"])
    assert body["chapter_count"] >= 1
    assert body["paragraph_count"] >= 1

    deleted = client.delete(f"/api/v1/books/{book_id}")
    assert deleted.status_code in {200, 204}, deleted.text
    assert client.get(f"/api/v1/books/{book_id}").status_code == 404
    assert fixture.is_file()
    assert fixture.read_bytes() == payload
